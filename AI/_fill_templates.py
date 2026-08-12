import random
import re
import sys
from pathlib import Path

import docx
from faker import Faker

fake = Faker("en_AU")

AU_STATES = ["NSW", "VIC", "QLD", "WA", "SA", "TAS", "ACT", "NT"]


def fake_name() -> str:
    return fake.name()


def fake_address() -> str:
    street = fake.street_address().replace("\n", ", ")
    return f"{street}, {fake.city()} {random.choice(AU_STATES)} {fake.postcode()}"


def fake_abn() -> str:
    return f"{random.randint(10, 99)} {random.randint(100, 999)} {random.randint(100, 999)} {random.randint(100, 999)}"


def fake_bsb() -> str:
    return f"{random.randint(100, 999)}-{random.randint(100, 999)}"


def fake_account_number() -> str:
    return f"{random.randint(1000, 9999)} {random.randint(1000, 9999)}"


def fake_dollar(lo: float = 10.0, hi: float = 900.0) -> str:
    return f"${random.uniform(lo, hi):.2f}"


def fake_dollar_thousands() -> str:
    return f"${random.randint(30, 150) * 1_000:,}"


def fake_hours() -> str:
    return f"{random.uniform(1.0, 40.0):.1f}"


def fake_date() -> str:
    d = fake.date_between(start_date="-5y", end_date="today")
    return d.strftime("%d/%m/%Y")


# field_key matches field_extraction_en.py's output dict keys (name,
# address, abn, bsb, account_number, salary) so a generated value can be
# checked against what extraction actually found. None means this filler
# doesn't correspond to any scored field (e.g. the employer's own name --
# only the employee's is scored).


def _apply_label_fillers(text: str, fillers: list[tuple[re.Pattern, object, str | None]]) -> tuple[str, list[tuple[str, str]]]:
    """Each fillers entry is (pattern, generator, field_key). pattern must
    have exactly one capture group covering the placeholder span to
    replace -- everything else in the line (the label text) is preserved
    verbatim. Applies every fillers entry that matches (a line can carry
    more than one field, e.g. the payslip's dense header cells)."""
    recorded: list[tuple[str, str]] = []
    for pattern, generator, field_key in fillers:
        m = pattern.search(text)
        if m is None:
            continue
        value = generator()
        text = text[: m.start(1)] + value + text[m.end(1):]
        if field_key is not None:
            recorded.append((field_key, value))
    return text, recorded


_DOLLAR_THOUSANDS_RE = re.compile(r"\$\d{2},\d{3}\b")
_DOLLAR_RE = re.compile(r"\$\d{2}\.\d{2}\b")
_HOURS_RE = re.compile(r"(?<!\$)\b\d{2}\.\d{2}\b")


def _apply_numeric_placeholders(text: str) -> tuple[str, list[str]]:
    """Blanket-replaces the payslip template's own bare '$00.00' / '$00,000'
    / '00.00' placeholder shapes with realistic random values, independent
    of any label -- these sit in table cells (hours/rate/total columns)
    with no adjacent label text to anchor on. Returns every dollar value
    substituted (each one is a legitimate 'salary' list-field ground-truth
    entry per SALARY_RE's shape) alongside the substituted text; hours
    values are generated but not returned since they aren't a scored
    field (SALARY_RE requires a '$' prefix)."""
    salary_values: list[str] = []

    def thousands_repl(_m: re.Match) -> str:
        value = fake_dollar_thousands()
        salary_values.append(value)
        return value

    def dollar_repl(_m: re.Match) -> str:
        value = fake_dollar()
        salary_values.append(value)
        return value

    def hours_repl(_m: re.Match) -> str:
        return fake_hours()

    text = _DOLLAR_THOUSANDS_RE.sub(thousands_repl, text)
    text = _DOLLAR_RE.sub(dollar_repl, text)
    text = _HOURS_RE.sub(hours_repl, text)
    return text, salary_values


def _set_paragraph_text(p, new_text: str) -> None:
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)


# Vetted against field_extraction_en.py's actual LABEL_PATTERNS -- of the 8
# downloaded en_contract templates with real placeholder blanks, only this
# one uses a colon-anchored "Employee Name:"/"Employee Address:" convention
# LABEL_PATTERNS actually matches; the rest embed placeholders in letter
# headers or prose ("[Employee First Name]", "Dear [insert employee's
# first name]") with no colon anchor, and were dropped rather than
# special-cased (see docs/superpowers/specs/2026-08-12-en-validation-
# corpus-benchmark-design.md, Decision 3).
CONTRACT_LABEL_FILLERS = [
    (re.compile(r"Employer\s*Name\s*:\s*(_{4,})", re.IGNORECASE), fake_name, None),
    (re.compile(r"Employer\s*Address\s*:\s*(_{4,})", re.IGNORECASE), fake_address, None),
    (re.compile(r"Employee\s*Name\s*:\s*(_{4,})", re.IGNORECASE), fake_name, "name"),
    (re.compile(r"Employee\s*Address\s*:\s*(_{4,})", re.IGNORECASE), fake_address, "address"),
]


def fill_docx_contract(src_path: str, dst_path: str, seed: int | None = None) -> tuple[int, dict[str, list[str]]]:
    if seed is not None:
        random.seed(seed)
        Faker.seed(seed)

    doc = docx.Document(src_path)
    filled = 0
    ground_truth: dict[str, list[str]] = {}

    for p in doc.paragraphs:
        text = p.text
        if not text.strip():
            continue
        new_text, recorded = _apply_label_fillers(text, CONTRACT_LABEL_FILLERS)
        if new_text != text:
            _set_paragraph_text(p, new_text)
            filled += 1
            for field_key, value in recorded:
                ground_truth.setdefault(field_key, []).append(value)

    Path(dst_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst_path)
    return filled, ground_truth


# The real Fair Work Ombudsman official payslip template
# (samples/en_pay_slip/Pay-slip-template-sts.docx) -- the same document
# field_extraction_en.py's docstring already cites as its grounding source.
# Structurally different from the contract template: fields are packed
# multiple-per-cell as separate paragraphs within one table cell (confirmed
# via direct inspection -- cell.paragraphs returns one paragraph per line,
# not one paragraph with embedded line breaks), using "<insert x>"
# placeholders, plus bare "$00.00"/"$00,000"/"00.00" numeric placeholders
# in the Entitlements/Deductions/Superannuation line-item tables with no
# adjacent label at all.
PAYSLIP_LABEL_FILLERS = [
    (re.compile(r"\*?Employee\s*:\s*(<[^<>]*>)", re.IGNORECASE), fake_name, "name"),
    (re.compile(r"\*?ABN\s*:\s*(<[^<>]*>)", re.IGNORECASE), fake_abn, "abn"),
    (re.compile(r"BSB\s*:\s*(<[^<>]*>)", re.IGNORECASE), fake_bsb, "bsb"),
    (re.compile(r"Account\s*:\s*(<[^<>]*>)", re.IGNORECASE), fake_account_number, "account_number"),
]


def fill_docx_payslip(src_path: str, dst_path: str, seed: int | None = None) -> tuple[int, dict[str, list[str]]]:
    if seed is not None:
        random.seed(seed)
        Faker.seed(seed)

    doc = docx.Document(src_path)
    filled = 0
    ground_truth: dict[str, list[str]] = {}

    def record(field_key: str, value: str) -> None:
        ground_truth.setdefault(field_key, []).append(value)

    for table in doc.tables:
        for row in table.rows:
            row_label = row.cells[0].text.strip().lower()
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    if not text.strip():
                        continue
                    new_text, recorded = _apply_label_fillers(text, PAYSLIP_LABEL_FILLERS)
                    new_text, salary_values = _apply_numeric_placeholders(new_text)
                    if new_text != text:
                        _set_paragraph_text(p, new_text)
                        filled += 1
                        for field_key, value in recorded:
                            record(field_key, value)
                        for value in salary_values:
                            record("salary", value)
            # "Total gross payment" (table 1, last row) is the only figure
            # extract_income_en() actually reads -- it checks gross before
            # net, so that's the value ground truth must match, even though
            # "Total net payment" (table 4) also gets a real dollar value
            # above via the generic pass (and correctly counts toward
            # "salary" list-field ground truth, matching real observed
            # production behavior where the gross total appears in both
            # the "salary" list and the "income" field simultaneously).
            if row_label.startswith("total gross payment"):
                m = re.search(r"\$[\d,]+\.\d{2}", row.cells[-1].text)
                if m:
                    ground_truth["income"] = [m.group(0)]
                    ground_truth["income_basis"] = ["gross"]

    Path(dst_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst_path)
    return filled, ground_truth


def fill_docx(src_path: str, dst_path: str, seed: int | None = None) -> tuple[int, dict[str, list[str]]]:
    if "en_pay_slip" in src_path.replace("\\", "/"):
        return fill_docx_payslip(src_path, dst_path, seed=seed)
    return fill_docx_contract(src_path, dst_path, seed=seed)


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2]
    seed = int(sys.argv[3]) if len(sys.argv) > 3 else None
    n, gt = fill_docx(src, dst, seed=seed)
    print(f"filled {n} fields: {dst}")
    print(f"ground truth: {gt}")
