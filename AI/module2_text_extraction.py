import csv
from pathlib import Path

import docx
import pypdfium2 as pdfium

from field_extraction_en import extract_fields_from_text_en


def extract_text_from_docx(path: str) -> str:
    text, _ = extract_docx_text_and_tables(path)
    return text


def extract_docx_text_and_tables(path: str) -> tuple[str, list[list[str]]]:
    doc = docx.Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    table_rows: list[list[str]] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            parts.append(" ".join(cells))
            table_rows.append(cells)
    return "\n".join(parts), table_rows


def extract_text_from_csv(path: str) -> str:
    with open(path, encoding="utf-8-sig", newline="") as f:
        return "\n".join(" ".join(row) for row in csv.reader(f))


def extract_text_from_pdf(path: str) -> str:
    # ponytail: pdfium keeps a native file handle open until this
    # PdfDocument is explicitly closed -- on Windows that blocks the
    # caller's os.unlink() of the tmp file (WinError 32), unlike POSIX
    # which allows deleting an open file. Close pages/doc explicitly
    # rather than relying on GC timing.
    pdf = pdfium.PdfDocument(path)
    try:
        pages = []
        for page in pdf:
            textpage = page.get_textpage()
            pages.append(textpage.get_text_range())
            textpage.close()
            page.close()
        return "\n".join(pages)
    finally:
        pdf.close()


_EXTRACTORS = {
    ".docx": extract_text_from_docx,
    ".csv": extract_text_from_csv,
    ".pdf": extract_text_from_pdf,
}


def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(f"unsupported text-native format: {ext}")
    return extractor(path)


def extract_fields(path: str) -> dict:
    if Path(path).suffix.lower() == ".docx":
        text, table_rows = extract_docx_text_and_tables(path)
        fields = extract_fields_from_text_en(text, table_rows=table_rows)
    else:
        text = extract_text(path)
        fields = extract_fields_from_text_en(text)
    return {"fields": fields, "text": text}


if __name__ == "__main__":
    import glob

    with open("module2_text_selfcheck_output.txt", "w", encoding="utf-8") as out:
        samples = glob.glob("samples/en_balance_sheet/*.docx") + glob.glob("samples/en_balance_sheet/*.pdf")
        samples += glob.glob("samples/en_contract/*.docx")
        for name in samples:
            result = extract_fields(name)
            out.write(f"--- {name} ---\n")
            for field, value in result["fields"].items():
                out.write(f"  {field}: {value}\n")
            out.write(f"  text length: {len(result['text'])} chars\n\n")
    print("saved module2_text_selfcheck_output.txt")
